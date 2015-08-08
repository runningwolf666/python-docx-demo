#!/usr/bin/env python3
#-*- coding:utf-8 -*-
请在Python3下运行此程序='Please run this program with Python3'

import time
import os

# Doc: http://python-docx.readthedocs.org/en/latest/index.html
# Quickstart: http://python-docx.readthedocs.org/en/latest/user/quickstart.html
# API Documentation: http://python-docx.readthedocs.org/en/latest/index.html#api-documentation
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Inches
from PIL import Image # Doc: http://pillow.readthedocs.org/

document = Document()

# 设置字体、大小
# run = document.add_paragraph().add_run()
# font = run.font
# font.name = 'Droid Sans'
# font.size = Pt(36)

# Each Document object provides access to its CoreProperties object via its core_properties attribute.
# About CoreProperties objects: https://python-docx.readthedocs.org/en/latest/api/document.html#id1
document.core_properties.author = 'author'
document.core_properties.comments = 'comments'

# document.add_heading('Document Title', 0)
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

# About RGBColor objects: https://python-docx.readthedocs.org/en/latest/api/shared.html#rgbcolor-objects
p.add_run('\n天蓝色').font.color.rgb = RGBColor(0x87, 0xCE, 0xEB) # 天蓝色
p.add_run('WHITE').font.color.rgb = RGBColor.from_string('FFFFFF') # white
p.add_run('RED').font.color.rgb = RGBColor.from_string('FF0000') # red
p.add_run('BLUE').font.color.rgb = RGBColor.from_string('0000FF') # blue
p.add_run('BLACK').font.color.rgb = RGBColor.from_string('000000') # black
# 紫色：800080    天蓝色：87CEEB    黄色：FFFF00
# 白色: FFFFFF    红色: FF0000    蓝色: 0000FF    黑色: 000000
# More: http://www.360doc.com/content/12/0219/22/19147_187921920.shtml

# add pic, original size
picname = 'pic.png'
if os.path.exists(picname):
    p.add_run('\nCode: ').font.color.rgb = RGBColor.from_string('0000FF') # blue
    # See: http://stackoverflow.com/questions/26617218/python-docx-add-picture-size-wont-print-as-shown-on-screen
    # document.add_picture('pic.png', width=Inches(4.9))
    im = Image.open(picname)
    width, height = im.size
    if width > 680:
        raise PicSizeError('Too large, please decrease the pic size or increase the number following--96')
    document.add_picture(picname, width=Inches(width/96))
    # document.add_picture(picname)

document.add_page_break()
document.add_paragraph().add_run('\nPowered by RunningWolf<runningwolf2016@163.com>').font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # white

timestyle = time.strftime('%H%M%S')
docxname = 'demo_{}.docx'.format(timestyle)
document.save(docxname)


class PicSizeError(Exception):
    """picture size error"""
    pass


